import multiprocessing
multiprocessing.freeze_support()

import os
import sys
import fitz  # PyMuPDF
import customtkinter as ctk
from tkinter import filedialog, messagebox
from tkinterdnd2 import DND_FILES, TkinterDnD
import concurrent.futures
import threading
import multiprocessing

import PIL.Image as PILImage
import PIL.ImageOps as ImageOps

# Intentar importar docx con reporte de error detallado
HAS_WORD = False
WORD_ERROR = ""
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    HAS_WORD = True
except Exception as e:
    HAS_WORD = False
    WORD_ERROR = str(e)
    print(f"Error cargando docx: {e}")

# =================================================================
# INTERFAZ DE USUARIO (GUI)
# =================================================================
from worker import convert_page_worker

class AIRISConverter(ctk.CTk, TkinterDnD.DnDWrapper):
    def __init__(self):
        # 1. Inicializamos CustomTkinter
        ctk.CTk.__init__(self)
        
        self.title("AIRIS Suite V1.1")
        self.geometry("600x550")
        self.resizable(False, False)
        
        # Colores Minimalistas
        self.bg_color = "#0F0F0F"
        self.accent_color = "#3B8ED0"
        self.card_color = "#1A1A1A"
        
        self.configure(fg_color=self.bg_color)
        
        # Estado de la Aplicación
        self.pdf_paths = []
        self.img_paths = []
        self.is_processing = False
        
        self._show_mode_selection()

    def _clear_window(self):
        for widget in self.winfo_children():
            widget.destroy()

    def _show_mode_selection(self):
        self._clear_window()
        
        # Header
        self.lbl_title = ctk.CTkLabel(
            self, text="AIRIS", 
            font=ctk.CTkFont(family="Inter", size=48, weight="bold"),
            text_color="white"
        )
        self.lbl_title.pack(pady=(60, 10))
        
        self.lbl_tag = ctk.CTkLabel(
            self, text="Selecciona el tipo de proceso", 
            font=ctk.CTkFont(size=14), text_color="#888888"
        )
        self.lbl_tag.pack(pady=(0, 40))

        # Botones de modo
        self.btn_full = ctk.CTkButton(
            self, text="NUEVO PROCESO COMPLETO\n(PDF → Imágenes → Word)", 
            command=self._setup_minimal_ui,
            height=80, width=350, corner_radius=15,
            fg_color=self.accent_color, hover_color="#2A6AA0",
            font=ctk.CTkFont(size=14, weight="bold")
        )
        self.btn_full.pack(pady=10)

        self.btn_only_word = ctk.CTkButton(
            self, text="SOLO GENERAR WORD\n(Imágenes → Word)", 
            command=self._setup_img_to_word_ui,
            height=80, width=350, corner_radius=15,
            fg_color="#222222", hover_color="#333333",
            font=ctk.CTkFont(size=14, weight="bold"),
            border_width=1, border_color="#444444"
        )
        self.btn_only_word.pack(pady=10)

        self.lbl_footer = ctk.CTkLabel(self, text="v1.1 | Desarrollado por Jorge Meneses", font=ctk.CTkFont(size=10), text_color="white")
        self.lbl_footer.pack(side="bottom", pady=15)

    def _setup_img_to_word_ui(self):
        self._clear_window()
        
        # Header simple
        self.btn_back = ctk.CTkButton(self, text="← Volver", width=80, height=30, fg_color="transparent", hover_color="#222222", command=self._show_mode_selection)
        self.btn_back.place(x=20, y=20)

        self.lbl_title = ctk.CTkLabel(self, text="AIRIS Solo Word", font=ctk.CTkFont(size=24, weight="bold"), text_color="white")
        self.lbl_title.pack(pady=(40, 10))

        # Drop Zone para imágenes
        self.drop_card = ctk.CTkFrame(self, fg_color=self.card_color, corner_radius=20, width=480, height=180)
        self.drop_card.pack(pady=10)
        self.drop_card.pack_propagate(False)
        
        self.lbl_drop = ctk.CTkLabel(
            self.drop_card, text="Suelta tus imágenes (JPG/PNG) aquí",
            font=ctk.CTkFont(size=14), text_color="#AAAAAA"
        )
        self.lbl_drop.place(relx=0.5, rely=0.5, anchor="center")

        # Permitir clic como fallback
        self.drop_card.bind("<Button-1>", lambda e: self._select_img_files())
        self.lbl_drop.bind("<Button-1>", lambda e: self._select_img_files())

        # Botones de selección manual
        self.sel_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.sel_frame.pack(pady=5)
        
        ctk.CTkButton(self.sel_frame, text="📁 Seleccionar Imágenes", width=160, command=self._select_img_files, fg_color="#333333").grid(row=0, column=0, padx=5)
        ctk.CTkButton(self.sel_frame, text="📂 Seleccionar Carpeta", width=160, command=self._select_img_folder, fg_color="#333333").grid(row=0, column=1, padx=5)

        # Botón Acción
        self.btn_main = ctk.CTkButton(
            self, text="GENERAR WORD", 
            command=self._start_only_word_process,
            height=50, width=300, corner_radius=25,
            fg_color="#27ae60", hover_color="#1e8449",
            font=ctk.CTkFont(size=15, weight="bold")
        )
        self.btn_main.pack(pady=40)

        self.lbl_footer = ctk.CTkLabel(self, text="v1.1 | Desarrollado por Jorge Meneses", font=ctk.CTkFont(size=10), text_color="white")
        self.lbl_footer.pack(side="bottom", pady=15)
        
        self._setup_dnd()

    def _setup_minimal_ui(self):
        self._clear_window()
        
        # Botón Volver
        self.btn_back = ctk.CTkButton(self, text="← Volver", width=80, height=30, fg_color="transparent", hover_color="#222222", command=self._show_mode_selection)
        self.btn_back.place(x=20, y=20)

        # Header
        self.lbl_title = ctk.CTkLabel(
            self, text="AIRIS", 
            font=ctk.CTkFont(family="Inter", size=32, weight="bold"),
            text_color="white"
        )
        self.lbl_title.pack(pady=(40, 5))
        
        self.lbl_tag = ctk.CTkLabel(
            self, text="Conversión Inteligente • V1.1", 
            font=ctk.CTkFont(size=12), text_color="#555555"
        )
        self.lbl_tag.pack(pady=(0, 30))

        # --- PASO 1: ZONA DE CARGA ---
        self.step1_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.step1_frame.pack(fill="x", padx=60)

        self.drop_card = ctk.CTkFrame(self.step1_frame, fg_color=self.card_color, corner_radius=20, width=480, height=180)
        self.drop_card.pack(pady=10)
        self.drop_card.pack_propagate(False)
        
        self.lbl_drop = ctk.CTkLabel(
            self.drop_card, text="PASO 1: Suelta tus PDF aquí",
            font=ctk.CTkFont(size=14), text_color="#AAAAAA"
        )
        self.lbl_drop.place(relx=0.5, rely=0.5, anchor="center")
        
        # Permitir clic también en la tarjeta como fallback
        self.drop_card.bind("<Button-1>", lambda e: self._select_files())
        self.lbl_drop.bind("<Button-1>", lambda e: self._select_files())

        # Botones de selección manual
        self.sel_frame_pdf = ctk.CTkFrame(self.step1_frame, fg_color="transparent")
        self.sel_frame_pdf.pack(pady=5)
        
        ctk.CTkButton(self.sel_frame_pdf, text="📄 Seleccionar Archivos", width=160, command=self._select_files, fg_color="#333333").grid(row=0, column=0, padx=5)
        ctk.CTkButton(self.sel_frame_pdf, text="📂 Seleccionar Carpeta", width=160, command=self._select_pdf_folder, fg_color="#333333").grid(row=0, column=1, padx=5)

        # --- PASO 2: CONFIGURACIÓN (Oculto inicialmente) ---
        self.step2_frame = ctk.CTkFrame(self, fg_color="transparent")
        # No hacemos pack() todavía

        self.lbl_step2 = ctk.CTkLabel(self.step2_frame, text="PASO 2: Ajustes de salida", font=ctk.CTkFont(size=13, weight="bold"), text_color=self.accent_color)
        self.lbl_step2.pack(pady=(20, 10))

        self.settings_inner = ctk.CTkFrame(self.step2_frame, fg_color="transparent")
        self.settings_inner.pack(fill="x")
        
        ctk.CTkLabel(self.settings_inner, text="Calidad:", font=ctk.CTkFont(size=12)).grid(row=0, column=0, sticky="w")
        self.opt_quality = ctk.CTkSegmentedButton(
            self.settings_inner, 
            values=["Estándar", "HD", "Ultra"],
            fg_color="#121212",
            selected_color=self.accent_color
        )
        self.opt_quality.set("HD")
        self.opt_quality.grid(row=0, column=1, padx=20, sticky="w")

        self.var_word = ctk.BooleanVar(value=True)
        self.chk_word = ctk.CTkCheckBox(
            self.step2_frame, text="Compilar automáticamente en Word (.docx)", 
            variable=self.var_word, border_width=2,
            font=ctk.CTkFont(size=12)
        )
        self.chk_word.pack(pady=20)

        self.btn_main = ctk.CTkButton(
            self.step2_frame, text="GENERAR DOCUMENTOS", 
            command=self._start_conversion_thread,
            height=50, width=300, corner_radius=25,
            fg_color=self.accent_color, hover_color="#2A6AA0",
            font=ctk.CTkFont(size=15, weight="bold")
        )
        self.btn_main.pack(pady=(0, 20))

        self.lbl_footer = ctk.CTkLabel(self, text="v1.1 | Desarrollado por Jorge Meneses", font=ctk.CTkFont(size=10), text_color="white")
        self.lbl_footer.pack(side="bottom", pady=15)

        # Registro DND
        self._setup_dnd()

    def _show_step_2(self):
        """Muestra la segunda parte de la interfaz"""
        if not self.step2_frame.winfo_ismapped():
            self.step2_frame.pack(fill="x", padx=60, after=self.step1_frame)
            # Cambiar texto del drop para indicar éxito
            self.lbl_drop.configure(text=f"✓ {len(self.pdf_paths)} Archivos Cargados", text_color=self.accent_color)
            self.drop_card.configure(border_width=2, border_color=self.accent_color)

    def _setup_dnd(self):
        """Registro robusto de Drag & Drop para Carpeta/Archivos"""
        try:
            target_widgets = [self.drop_card, self.lbl_drop]
            for widget in target_widgets:
                try:
                    # Intentar via wrapper de tkinterdnd2
                    widget.drop_target_register(DND_FILES)
                    widget.dnd_bind('<<Drop>>', self._on_drop)
                except:
                    # Fallback a Tcl directo si el wrapper no está disponible
                    try:
                        self.tk.call('tkdnd::drop_target', 'register', widget._w, (DND_FILES,))
                        # Puente para recibir el dato raw
                        def _bridge(data):
                            class FakeEvent: pass
                            e = FakeEvent()
                            e.data = data
                            self._on_drop(e)
                        tcl_cmd = self.register(_bridge)
                        self.tk.call('bind', widget._w, '<<Drop>>', f'{tcl_cmd} %D')
                    except: pass
        except: pass

    def _on_drop(self, event):
        """Procesa archivos y carpetas soltados"""
        try:
            # TkinterDnD2 puede devolver una cadena o una tupla
            if isinstance(event.data, (list, tuple)):
                files = event.data
            else:
                raw_data = event.data
                # Manejar el formato de Tcl {C:/Path with spaces}
                if raw_data.startswith('{') and raw_data.endswith('}'):
                    raw_data = raw_data[1:-1]
                files = self.tk.splitlist(raw_data)
        except Exception as e:
            print(f"Error parseando drop: {e}")
            return

        extracted_files = []
        
        # Modo PDF o Imagen según UI actual
        is_pdf_mode = hasattr(self, "lbl_step2")
        valid_exts = (".pdf",) if is_pdf_mode else (".jpg", ".jpeg", ".png")

        for item in files:
            # Normalizar ruta
            item = item.strip('{}') 
            if os.path.exists(item):
                if os.path.isdir(item):
                    for root, _, filenames in os.walk(item):
                        for f in filenames:
                            if f.lower().endswith(valid_exts):
                                extracted_files.append(os.path.join(root, f))
                elif item.lower().endswith(valid_exts):
                    extracted_files.append(item)

        if extracted_files:
            if is_pdf_mode:
                self.pdf_paths.extend(extracted_files)
                self._show_step_2()
            else:
                self.img_paths.extend(extracted_files)
                self.lbl_drop.configure(text=f"✓ {len(self.img_paths)} Imágenes Listas", text_color="#27ae60")
                self.drop_card.configure(border_width=2, border_color="#27ae60")

    def _select_img_files(self):
        files = filedialog.askopenfilenames(filetypes=[("Imágenes", "*.jpg *.jpeg *.png")])
        if files:
            self.img_paths.extend(list(files))
            self.lbl_drop.configure(text=f"✓ {len(self.img_paths)} Imágenes Listas", text_color="#27ae60")
            self.drop_card.configure(border_width=2, border_color="#27ae60")

    def _select_img_folder(self):
        folder = filedialog.askdirectory(title="Selecciona carpeta con imágenes")
        if folder:
            valid_exts = ('.jpg', '.jpeg', '.png')
            found = [os.path.join(folder, f) for f in os.listdir(folder) if f.lower().endswith(valid_exts)]
            if found:
                self.img_paths.extend(found)
                self.lbl_drop.configure(text=f"✓ {len(self.img_paths)} Imágenes Listas", text_color="#27ae60")
                self.drop_card.configure(border_width=2, border_color="#27ae60")
            else:
                messagebox.showwarning("Atención", "No se encontraron imágenes JPG/PNG en esa carpeta.")

    def _select_pdf_folder(self):
        folder = filedialog.askdirectory(title="Selecciona carpeta con PDFs")
        if folder:
            found = [os.path.join(folder, f) for f in os.listdir(folder) if f.lower().endswith(".pdf")]
            if found:
                self.pdf_paths.extend(found)
                self._show_step_2()
            else:
                messagebox.showwarning("Atención", "No se encontraron archivos PDF en esa carpeta.")

    def _start_only_word_process(self):
        if not self.img_paths:
            messagebox.showwarning("Atención", "Selecciona o arrastra imágenes primero.")
            return
        
        desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        out_target = os.path.join(desktop, "AIRIS_SOLO_WORD")
        if not os.path.exists(out_target): os.makedirs(out_target)
        
        self.is_processing = True
        self.btn_main.configure(state="disabled", text="GENERANDO...")
        
        self._show_progress_window()
        self.lbl_proc.configure(text="Compilando imágenes...")
        
        def run():
            try:
                self._process_word_generation_from_list(self.img_paths, out_target)
                self._finish_all(f"¡Documento generado!\n\nUbicación: {out_target}")
                try: os.startfile(out_target)
                except: pass
            except Exception as e:
                self._finish_all(f"Error: {e}")

        threading.Thread(target=run, daemon=True).start()

    def _process_word_generation_from_list(self, img_list, save_dir):
        try:
            img_list.sort()
            doc = Document()
            total = len(img_list)
            for i, img_path in enumerate(img_list):
                self.prog_bar.set((i+1)/total)
                if i > 0: section = doc.add_section()
                else: section = doc.sections[0]
                
                try:
                    with PILImage.open(img_path) as img: w_px, h_px = img.size
                except: continue
                
                section.orientation = WD_ORIENT.PORTRAIT
                section.page_width = Inches(8.5); section.page_height = Inches(11)
                available_w = 7.5; available_h = 10.0
                
                img_w_in = w_px / 96; img_h_in = h_px / 96
                scale = 1.0
                if img_w_in > available_w: scale = min(scale, available_h / img_h_in, available_w / img_w_in)
                
                final_w = img_w_in * scale; final_h = img_h_in * scale
                section.top_margin = Inches((11.0 - final_h) / 2)
                section.left_margin = section.right_margin = Inches(0.5)
                
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(final_w))
            
            doc.save(os.path.join(save_dir, "AIRIS_Solo_Word.docx"))
        except Exception as e:
            print(f"Error Word: {e}")
            raise e

    def _select_files(self):
        files = filedialog.askopenfilenames(filetypes=[("PDF Files", "*.pdf")])
        if files:
            self.pdf_paths.extend(list(files))
            self._show_step_2()

    def _show_progress_window(self):
        """Crea una ventana flotante para el progreso"""
        self.progress_win = ctk.CTkToplevel(self)
        self.progress_win.title("Procesando...")
        self.progress_win.geometry("400x150")
        self.progress_win.resizable(False, False)
        self.progress_win.attributes("-topmost", True)
        self.progress_win.configure(fg_color=self.bg_color)
        
        # Centrar respecto a la principal
        self.update_idletasks()
        x = self.winfo_x() + (self.winfo_width() // 2) - 200
        y = self.winfo_y() + (self.winfo_height() // 2) - 75
        self.progress_win.geometry(f"+{x}+{y}")

        self.lbl_proc = ctk.CTkLabel(self.progress_win, text="Iniciando conversión...", font=ctk.CTkFont(weight="bold"))
        self.lbl_proc.pack(pady=(30, 10))

        self.prog_bar = ctk.CTkProgressBar(self.progress_win, width=320)
        self.prog_bar.set(0)
        self.prog_bar.pack(pady=10)

    def _start_conversion_thread(self):
        if not self.pdf_paths:
            messagebox.showwarning("Atención", "Suelta o selecciona algunos archivos PDF primero.")
            return

        # Ruta default: Escritorio/AIRIS_SALIDA
        desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        out_target = os.path.join(desktop, "AIRIS_SALIDA")
        if not os.path.exists(out_target): os.makedirs(out_target)
        
        self.is_processing = True
        self.btn_main.configure(state="disabled", text="CONVIRTIENDO...")
        
        self._show_progress_window()
        
        thread = threading.Thread(target=self._process_conversion, args=(out_target,))
        thread.daemon = True
        thread.start()

    def _process_conversion(self, output_base):
        try:
            # Mapeo de calidad
            q_choice = self.opt_quality.get()
            dpi = 300
            if q_choice == "Estándar": dpi = 150
            elif q_choice == "Ultra": dpi = 600
            
            output_folder = os.path.join(output_base, "IMAGENES")
            if not os.path.exists(output_folder): os.makedirs(output_folder)

            tasks = []
            for pdf_path in self.pdf_paths:
                try:
                    pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
                    doc = fitz.open(pdf_path)
                    for i in range(len(doc)):
                        out_path = os.path.join(output_folder, f"{pdf_name}_p{i+1}.jpg")
                        tasks.append((pdf_path, i, out_path, dpi, 95, True))
                    doc.close()
                except: pass

            total = len(tasks)
            if total == 0:
                self._finish_all("No hay contenido para procesar.")
                return

            num_cores = multiprocessing.cpu_count()
            completed = 0
            with concurrent.futures.ProcessPoolExecutor(max_workers=num_cores) as executor:
                futures = [executor.submit(convert_page_worker, task) for task in tasks]
                for future in concurrent.futures.as_completed(futures):
                    completed += 1
                    self.prog_bar.set(completed / total)
                    self.lbl_proc.configure(text=f"Procesando: {completed} / {total} páginas")
            
            # Word
            if self.var_word.get() and HAS_WORD:
                self.lbl_proc.configure(text="Generando Word...")
                word_folder = os.path.join(output_base, "DOCUMENTOS")
                if not os.path.exists(word_folder): os.makedirs(word_folder)
                self._process_word_generation(output_folder, word_folder)
            
            self._finish_all(f"¡Proceso completado!\n\nArchivos en: {output_base}")
            try: os.startfile(output_base)
            except: pass

        except Exception as e:
            self._finish_all(f"Error: {e}")

    def _process_word_generation(self, img_dir, save_dir):
        try:
            valid_exts = ('.jpg', '.jpeg', '.png')
            images = [os.path.join(img_dir, f) for f in os.listdir(img_dir) if f.lower().endswith(valid_exts)]
            images.sort()
            if not images: return

            doc = Document()
            for i, img_path in enumerate(images):
                if i > 0: section = doc.add_section()
                else: section = doc.sections[0]
                
                try:
                    with PILImage.open(img_path) as img: w_px, h_px = img.size
                except: continue
                
                section.orientation = WD_ORIENT.PORTRAIT
                section.page_width = Inches(8.5); section.page_height = Inches(11)
                available_w = 7.5; available_h = 10.0
                
                img_w_in = w_px / 96; img_h_in = h_px / 96
                scale = 1.0
                if img_w_in > available_w: scale = min(scale, available_h / img_h_in, available_w / img_w_in)
                
                final_w = img_w_in * scale; final_h = img_h_in * scale
                section.top_margin = Inches((11.0 - final_h) / 2)
                section.left_margin = section.right_margin = Inches(0.5)
                
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(final_w))
            
            doc.save(os.path.join(save_dir, "AIRIS_Compilado.docx"))
        except: pass

    def _finish_all(self, msg):
        self.is_processing = False
        self.pdf_paths = []
        self.img_paths = []
        self.btn_main.configure(state="normal", text="GENERAR DOCUMENTOS")
        if hasattr(self, 'progress_win'): self.progress_win.destroy()
        self.lbl_drop.configure(text="Arrastra archivos aquí\no haz clic para buscar", text_color="#AAAAAA")
        if msg: messagebox.showinfo("AIRIS", msg)

# =================================================================
# PUNTO DE ENTRADA
# =================================================================
if __name__ == "__main__":
    multiprocessing.freeze_support()
    app = AIRISConverter()
    app.mainloop()
